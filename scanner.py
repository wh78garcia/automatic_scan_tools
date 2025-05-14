import requests
from bs4 import BeautifulSoup
import argparse
from docx import Document
from urllib.parse import urljoin
import os
from concurrent.futures import ThreadPoolExecutor

class VulnScanner:
    def __init__(self, target_url):
        """Initialize the vulnerability scanner with target URL"""
        self.target_url = target_url
        self.session = requests.Session()
        self.session.headers = {"User-Agent": "Mozilla/5.0"}  # Set default user agent
        self.vulnerabilities = []  # Store found vulnerabilities

    def load_wordlist(self, file_path):
        """Load wordlist file for directory/payload testing"""
        with open(file_path, 'r') as f:
            return [line.strip() for line in f if line.strip()]

    def scan_directories(self):
        """Scan for sensitive directories using wordlist"""
        wordlist = self.load_wordlist("wordlists/dirs.txt")
        for dir_path in wordlist:
            full_url = urljoin(self.target_url, dir_path)
            
            try:
                response = self.session.get(full_url, timeout=5)
                if response.status_code == 200:
                    # print('[DEBUG] Testing URL:', full_url)
                    # print('[DEBUG] Response:', response.status_code)
                    self.vulnerabilities.append({
                        "type": "Exposed Directory",
                        "url": full_url,
                        "payload": None
                    })
                    print(f"[+] Found exposed directory: {full_url}")
            except requests.exceptions.RequestException:
                continue

    def test_sql_injection(self):
        """Test for SQL injection vulnerabilities"""
        test_url = urljoin(self.target_url, "/user?id=")
        payloads = self.load_wordlist("wordlists/payloads.txt")

        for payload in payloads:
            test_url_sql = test_url + payload
            
            try:
                response = self.session.get(test_url_sql, timeout=5)
                # Check if response contains evidence of successful injection
                if "admin" in response.text.lower() and "guest" in response.text.lower():
                    # print('[DEBUG] Testing SQLi payload:', test_url_sql)
                    # print('[DEBUG] Response:', response.text)  # Print first 50 chars
                    self.vulnerabilities.append({
                        "type": "SQL Injection",
                        "url": test_url_sql,
                        "payload": payload
                    })
                    print(f"[+] Potential SQL Injection found: {test_url_sql}")
            except requests.exceptions.RequestException:
                continue

    def generate_report(self):
        """Generate Word document vulnerability report"""
        doc = Document()
        doc.add_heading(f'Vulnerability Scan Report - {self.target_url}', level=1)
        doc.add_paragraph(f"Target URL: {self.target_url}\n")

        if not self.vulnerabilities:
            doc.add_paragraph("No critical vulnerabilities found.")
        else:
            # Create vulnerabilities table
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Shading"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Vulnerability Type"
            hdr_cells[1].text = "URL"
            hdr_cells[2].text = "Payload"

            # Add each vulnerability to the table
            for vuln in self.vulnerabilities:
                row_cells = table.add_row().cells
                row_cells[0].text = vuln["type"]
                row_cells[1].text = vuln["url"]
                row_cells[2].text = str(vuln["payload"])

        # Save report with sanitized filename
        report_path = f"reports/report_{self.target_url.replace('://', '_').replace('/', '_')}.docx"
        doc.save(report_path)
        print(f"\n[+] Report generated: {os.path.abspath(report_path)}")

if __name__ == "__main__":
    # Configure command line arguments
    parser = argparse.ArgumentParser(description="Automated Vulnerability Scanner")
    parser.add_argument("-u", "--url", 
                      default='http://127.0.0.1:5000', 
                      required=True, 
                      help="Target URL")
    args = parser.parse_args()

    # Create reports directory if not exists
    if not os.path.exists("reports"):
        os.makedirs("reports")

    # Initialize scanner
    scanner = VulnScanner(args.url)
    print(f"[*] Starting scan: {args.url}")

    # Run scans concurrently
    with ThreadPoolExecutor(max_workers=3) as executor:
        executor.submit(scanner.scan_directories)  # Directory scanning
        executor.submit(scanner.test_sql_injection)  # SQL injection testing

    # Generate final report
    scanner.generate_report()